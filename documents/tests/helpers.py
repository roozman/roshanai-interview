from io import BytesIO
from tempfile import TemporaryDirectory

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from docx import Document as WordDocument
from documents.constants import EMBEDDING_DIMENSION


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def build_docx_bytes(
    paragraphs: tuple[str, ...] = (),
    table_rows: tuple[tuple[str, ...], ...] = (),
) -> bytes:
    document = WordDocument()

    for text in paragraphs:
        document.add_paragraph(text)

    if table_rows:
        column_count = max(len(row) for row in table_rows)
        table = document.add_table(rows=0, cols=column_count)

        for values in table_rows:
            cells = table.add_row().cells

            for index, value in enumerate(values):
                cells[index].text = value

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()


def make_uploaded_docx(
    name: str = "sample.docx",
    paragraphs: tuple[str, ...] = (),
    table_rows: tuple[tuple[str, ...], ...] = (),
) -> SimpleUploadedFile:
    return SimpleUploadedFile(
        name=name,
        content=build_docx_bytes(
            paragraphs=paragraphs,
            table_rows=table_rows,
        ),
        content_type=DOCX_CONTENT_TYPE,
    )


class TemporaryMediaRootMixin:
    @classmethod
    def setUpClass(cls):
        cls.temporary_media = TemporaryDirectory()
        cls.settings_override = override_settings(
            MEDIA_ROOT=cls.temporary_media.name
        )
        cls.settings_override.enable()

        super().setUpClass()

    @classmethod
    def tearDownClass(cls):
        try:
            super().tearDownClass()
        finally:
            cls.settings_override.disable()
            cls.temporary_media.cleanup()

def build_test_embeddings(
    texts,
) -> list[list[float]]:
    return [
        [0.1] * EMBEDDING_DIMENSION
        for _ in texts
    ]
